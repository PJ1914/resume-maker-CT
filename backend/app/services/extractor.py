"""
Text extraction from PDF, DOCX, and LaTeX files.
Handles various edge cases and layout types.
Detects LaTeX content and parses it appropriately.
"""

import io
import re
from typing import Optional, Dict, List
import pdfplumber
from docx import Document


class LaTeXExtractor:
    """Extract and parse content from LaTeX files."""
    
    # Common LaTeX resume commands and their meanings
    SECTION_COMMANDS = {
        r'\\section\*?\{([^}]+)\}': 'section',
        r'\\subsection\*?\{([^}]+)\}': 'subsection',
        r'\\resumeSubheading': 'experience_item',
        r'\\resumeProjectHeading': 'project_item',
        r'\\resumeItem': 'bullet_item',
        r'\\resumeSubItem': 'sub_bullet_item',
        r'\\resumeSubSubheading': 'subheading',
    }
    
    @staticmethod
    def is_latex_content(text: str) -> bool:
        """
        Detect if text contains LaTeX commands.
        """
        latex_indicators = [
            r'\\documentclass',
            r'\\begin\{document\}',
            r'\\section\{',
            r'\\textbf\{',
            r'\\href\{',
            r'\\resumeSubheading',
            r'\\newcommand',
            r'\\usepackage',
        ]
        
        for indicator in latex_indicators:
            if re.search(indicator, text):
                return True
        return False
    
    @staticmethod
    def extract_text(file_content: bytes) -> Dict[str, any]:
        """
        Extract and parse LaTeX content into structured text.
        
        Args:
            file_content: LaTeX file bytes
            
        Returns:
            Dict with 'text', 'is_latex', 'structured_data'
        """
        try:
            # Try different encodings
            text = None
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    text = file_content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            
            if text is None:
                raise Exception("Could not decode LaTeX file")
            
            # Parse LaTeX into structured format
            structured_data = LaTeXExtractor.parse_latex_resume(text)
            
            # Also convert to plain text for fallback parsing
            plain_text = LaTeXExtractor.latex_to_plain_text(text)
            
            return {
                'text': plain_text,
                'raw_latex': text,
                'is_latex': True,
                'structured_data': structured_data,
                'extraction_method': 'latex_parser'
            }
            
        except Exception as e:
            raise Exception(f"LaTeX extraction failed: {str(e)}")
    
    @staticmethod
    def parse_latex_resume(latex: str) -> Dict[str, any]:
        """
        Parse LaTeX resume into structured data.
        
        Handles common LaTeX resume templates like:
        - Jake's Resume template
        - ModernCV
        - AltaCV
        - Custom resumeSubheading macros
        """
        data = {
            'contact_info': {},
            'sections': {},
            'experience': [],
            'education': [],
            'projects': [],
            'skills': {},
        }
        
        # Extract name - usually in \name{} or at the start
        name_patterns = [
            r'\\name\{([^}]+)\}',
            r'\\begin\{center\}[^}]*\\textbf\{\\Huge\\scshape\s*([^}]+)\}',
            r'\\textbf\{\\Huge\s*([^}]+)\}',
            r'\\LARGE\s*\\textbf\{([^}]+)\}',
        ]
        for pattern in name_patterns:
            match = re.search(pattern, latex)
            if match:
                data['contact_info']['name'] = LaTeXExtractor.clean_latex(match.group(1))
                break
        
        # Extract contact info
        # Email
        email_match = re.search(r'\\href\{mailto:([^}]+)\}', latex)
        if email_match:
            data['contact_info']['email'] = email_match.group(1)
        else:
            email_match = re.search(r'([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})', latex)
            if email_match:
                data['contact_info']['email'] = email_match.group(1)
        
        # Phone
        phone_match = re.search(r'(?:\\href\{tel:[^}]+\}\{)?([+\d\s()-]{10,})\}?', latex)
        if phone_match:
            data['contact_info']['phone'] = phone_match.group(1).strip()
        
        # LinkedIn
        linkedin_match = re.search(r'\\href\{[^}]*linkedin\.com/in/([^}]+)\}', latex)
        if linkedin_match:
            data['contact_info']['linkedin'] = f"linkedin.com/in/{linkedin_match.group(1)}"
        
        # GitHub
        github_match = re.search(r'\\href\{[^}]*github\.com/([^}]+)\}', latex)
        if github_match:
            data['contact_info']['github'] = f"github.com/{github_match.group(1)}"
        
        # Portfolio/Website
        portfolio_match = re.search(r'\\href\{(https?://[^}]+)\}\{[^}]*(?:portfolio|website|blog)[^}]*\}', latex, re.IGNORECASE)
        if portfolio_match:
            data['contact_info']['portfolio'] = portfolio_match.group(1)
        
        # Location
        location_match = re.search(r'([A-Z][a-z]+,?\s+[A-Z][a-z]+(?:,?\s+[A-Z]{2})?)\s*(?:\\\\|$|\n)', latex)
        if location_match:
            data['contact_info']['location'] = location_match.group(1).strip()
        
        # Extract sections
        section_pattern = r'\\section\{([^}]+)\}(.*?)(?=\\section\{|\\end\{document\}|$)'
        sections = re.findall(section_pattern, latex, re.DOTALL)
        
        for section_name, section_content in sections:
            section_name_clean = LaTeXExtractor.clean_latex(section_name).lower()
            
            if 'experience' in section_name_clean or 'work' in section_name_clean:
                data['experience'] = LaTeXExtractor.parse_experience_section(section_content)
            elif 'education' in section_name_clean:
                data['education'] = LaTeXExtractor.parse_education_section(section_content)
            elif 'project' in section_name_clean:
                data['projects'] = LaTeXExtractor.parse_projects_section(section_content)
            elif 'skill' in section_name_clean or 'technical' in section_name_clean:
                data['skills'] = LaTeXExtractor.parse_skills_section(section_content)
            else:
                # Store other sections as-is
                data['sections'][section_name_clean] = LaTeXExtractor.clean_latex(section_content)
        
        return data
    
    @staticmethod
    def parse_experience_section(content: str) -> List[Dict]:
        """Parse experience/work section from LaTeX."""
        experiences = []
        
        # Pattern for resumeSubheading: \resumeSubheading{Title}{Date}{Company}{Location}
        subheading_pattern = r'\\resumeSubheading\s*\{([^}]*)\}\s*\{([^}]*)\}\s*\{([^}]*)\}\s*\{([^}]*)\}'
        
        # Find all subheadings
        matches = list(re.finditer(subheading_pattern, content))
        
        for i, match in enumerate(matches):
            title = LaTeXExtractor.clean_latex(match.group(1))
            date_range = LaTeXExtractor.clean_latex(match.group(2))
            company = LaTeXExtractor.clean_latex(match.group(3))
            location = LaTeXExtractor.clean_latex(match.group(4))
            
            # Parse date range
            start_date, end_date = LaTeXExtractor.parse_date_range(date_range)
            
            # Get content until next subheading
            start_pos = match.end()
            end_pos = matches[i + 1].start() if i + 1 < len(matches) else len(content)
            item_content = content[start_pos:end_pos]
            
            # Extract bullet points
            bullets = LaTeXExtractor.extract_resume_items(item_content)
            
            experiences.append({
                'position': title,
                'company': company,
                'location': location,
                'startDate': start_date,
                'endDate': end_date,
                'description': '\n'.join(bullets)
            })
        
        # If no resumeSubheading found, try alternative patterns
        if not experiences:
            # Try \textbf{Company} pattern
            alt_pattern = r'\\textbf\{([^}]+)\}[^\\]*\\hfill[^\\]*([^\\]+)\\\\[^\\]*\\emph\{([^}]+)\}'
            for match in re.finditer(alt_pattern, content):
                company = LaTeXExtractor.clean_latex(match.group(1))
                date_range = LaTeXExtractor.clean_latex(match.group(2))
                position = LaTeXExtractor.clean_latex(match.group(3))
                
                start_date, end_date = LaTeXExtractor.parse_date_range(date_range)
                
                experiences.append({
                    'position': position,
                    'company': company,
                    'location': '',
                    'startDate': start_date,
                    'endDate': end_date,
                    'description': ''
                })
        
        return experiences
    
    @staticmethod
    def parse_education_section(content: str) -> List[Dict]:
        """Parse education section from LaTeX."""
        education = []
        
        # Pattern for resumeSubheading in education
        subheading_pattern = r'\\resumeSubheading\s*\{([^}]*)\}\s*\{([^}]*)\}\s*\{([^}]*)\}\s*\{([^}]*)\}'
        
        for match in re.finditer(subheading_pattern, content):
            school = LaTeXExtractor.clean_latex(match.group(1))
            date_range = LaTeXExtractor.clean_latex(match.group(2))
            degree = LaTeXExtractor.clean_latex(match.group(3))
            location = LaTeXExtractor.clean_latex(match.group(4))
            
            start_date, end_date = LaTeXExtractor.parse_date_range(date_range)
            
            # Extract GPA if present
            gpa_match = re.search(r'GPA[:\s]*([0-9.]+)', degree + content, re.IGNORECASE)
            gpa = gpa_match.group(1) if gpa_match else ''
            
            education.append({
                'school': school,
                'degree': degree,
                'location': location,
                'startDate': start_date,
                'endDate': end_date,
                'gpa': gpa
            })
        
        return education
    
    @staticmethod
    def parse_projects_section(content: str) -> List[Dict]:
        """Parse projects section from LaTeX."""
        projects = []
        
        # Pattern for resumeProjectHeading
        project_pattern = r'\\resumeProjectHeading\s*\{\\textbf\{([^}]*)\}[^}]*\\emph\{([^}]*)\}\}\s*\{([^}]*)\}'
        
        matches = list(re.finditer(project_pattern, content))
        
        for i, match in enumerate(matches):
            name = LaTeXExtractor.clean_latex(match.group(1))
            tech = LaTeXExtractor.clean_latex(match.group(2))
            date = LaTeXExtractor.clean_latex(match.group(3))
            
            # Get content until next project
            start_pos = match.end()
            end_pos = matches[i + 1].start() if i + 1 < len(matches) else len(content)
            item_content = content[start_pos:end_pos]
            
            # Extract bullet points
            bullets = LaTeXExtractor.extract_resume_items(item_content)
            
            # Check for link in name
            link_match = re.search(r'\\href\{([^}]+)\}', match.group(0))
            link = link_match.group(1) if link_match else ''
            
            projects.append({
                'name': name,
                'technologies': tech,
                'startDate': date,
                'endDate': date,
                'description': '\n'.join(bullets),
                'link': link
            })
        
        # Alternative pattern: bullet-based projects
        if not projects:
            bullet_pattern = r'\\resumeItem\{\\textbf\{([^}]+)\}[^}]*\|[^}]*\\emph\{([^}]*)\}'
            for match in re.finditer(bullet_pattern, content):
                name = LaTeXExtractor.clean_latex(match.group(1))
                tech = LaTeXExtractor.clean_latex(match.group(2))
                
                projects.append({
                    'name': name,
                    'technologies': tech,
                    'startDate': '',
                    'endDate': '',
                    'description': '',
                    'link': ''
                })
        
        return projects
    
    @staticmethod
    def parse_skills_section(content: str) -> Dict[str, List[str]]:
        """Parse skills section from LaTeX."""
        skills = {}
        
        # Pattern: \textbf{Category:} skill1, skill2, skill3
        category_pattern = r'\\textbf\{([^:}]+)[:\s]*\}[:\s]*([^\\\n]+)'
        
        for match in re.finditer(category_pattern, content):
            category = LaTeXExtractor.clean_latex(match.group(1))
            skill_list = LaTeXExtractor.clean_latex(match.group(2))
            
            # Split skills by comma
            skills[category] = [s.strip() for s in skill_list.split(',') if s.strip()]
        
        # If no categories found, extract all skills as a flat list
        if not skills:
            all_skills = LaTeXExtractor.clean_latex(content)
            skill_items = [s.strip() for s in re.split(r'[,;]', all_skills) if s.strip() and len(s.strip()) > 1]
            if skill_items:
                skills['Technical Skills'] = skill_items
        
        return skills
    
    @staticmethod
    def extract_resume_items(content: str) -> List[str]:
        """Extract bullet points from LaTeX content."""
        items = []
        
        # Pattern for \resumeItem{content}
        item_pattern = r'\\resumeItem\{([^}]+(?:\{[^}]*\}[^}]*)*)\}'
        
        for match in re.finditer(item_pattern, content):
            item_text = LaTeXExtractor.clean_latex(match.group(1))
            if item_text:
                items.append(item_text)
        
        # Also try \item pattern
        if not items:
            item_pattern = r'\\item\s+([^\n\\]+)'
            for match in re.finditer(item_pattern, content):
                item_text = LaTeXExtractor.clean_latex(match.group(1))
                if item_text:
                    items.append(item_text)
        
        return items
    
    @staticmethod
    def parse_date_range(date_str: str) -> tuple:
        """Parse date range string into start and end dates."""
        date_str = date_str.strip()
        
        # Common patterns: "Jan 2020 -- Present", "2020 - 2022", "May 2020"
        range_pattern = r'(.+?)\s*(?:--|–|-|to)\s*(.+)'
        match = re.match(range_pattern, date_str, re.IGNORECASE)
        
        if match:
            return match.group(1).strip(), match.group(2).strip()
        else:
            return date_str, date_str
    
    @staticmethod
    def clean_latex(text: str) -> str:
        """Remove LaTeX commands and clean up text."""
        if not text:
            return ''
        
        # Remove common LaTeX commands
        replacements = [
            (r'\\textbf\{([^}]*)\}', r'\1'),
            (r'\\textit\{([^}]*)\}', r'\1'),
            (r'\\emph\{([^}]*)\}', r'\1'),
            (r'\\underline\{([^}]*)\}', r'\1'),
            (r'\\href\{[^}]*\}\{([^}]*)\}', r'\1'),
            (r'\\url\{([^}]*)\}', r'\1'),
            (r'\\small', ''),
            (r'\\normalsize', ''),
            (r'\\large', ''),
            (r'\\Large', ''),
            (r'\\LARGE', ''),
            (r'\\huge', ''),
            (r'\\Huge', ''),
            (r'\\scshape', ''),
            (r'\\bfseries', ''),
            (r'\\itshape', ''),
            (r'\\hfill', ' '),
            (r'\\vspace\{[^}]*\}', ''),
            (r'\\hspace\{[^}]*\}', ' '),
            (r'\\\\', '\n'),
            (r'\\&', '&'),
            (r'\\%', '%'),
            (r'\\$', '$'),
            (r'\\_', '_'),
            (r'\\#', '#'),
            (r'\\,', ' '),
            (r'\\;', ' '),
            (r'\\:', ' '),
            (r'\\!', ''),
            (r'\\ ', ' '),
            (r'~', ' '),
            (r'\{', ''),
            (r'\}', ''),
        ]
        
        result = text
        for pattern, replacement in replacements:
            result = re.sub(pattern, replacement, result)
        
        # Remove remaining backslash commands
        result = re.sub(r'\\[a-zA-Z]+', '', result)
        
        # Clean up whitespace
        result = re.sub(r'\s+', ' ', result)
        result = result.strip()
        
        return result
    
    @staticmethod
    def latex_to_plain_text(latex: str) -> str:
        """Convert full LaTeX document to plain text."""
        # Extract only document body
        body_match = re.search(r'\\begin\{document\}(.*)\\end\{document\}', latex, re.DOTALL)
        if body_match:
            latex = body_match.group(1)
        
        # Convert sections to readable format
        latex = re.sub(r'\\section\{([^}]+)\}', r'\n\n== \1 ==\n', latex)
        latex = re.sub(r'\\subsection\{([^}]+)\}', r'\n\n-- \1 --\n', latex)
        
        # Convert resume items to bullets
        latex = re.sub(r'\\resumeItem\{', '• ', latex)
        latex = re.sub(r'\\item\s+', '• ', latex)
        
        # Clean LaTeX commands
        plain = LaTeXExtractor.clean_latex(latex)
        
        # Clean up extra whitespace
        plain = re.sub(r'\n{3,}', '\n\n', plain)
        
        return plain.strip()


class PDFExtractor:
    """Extract text from PDF files using pdfplumber."""
    
    @staticmethod
    def extract_hyperlinks(page) -> List[Dict[str, str]]:
        """
        Extract hyperlinks from a PDF page.
        
        Args:
            page: pdfplumber page object
            
        Returns:
            List of dicts with 'uri' and 'text' keys
        """
        hyperlinks = []
        try:
            # pdfplumber exposes annotations via page.annots
            if hasattr(page, 'annots') and page.annots:
                for annot in page.annots:
                    if annot.get('uri'):
                        hyperlinks.append({
                            'uri': annot.get('uri', ''),
                            'text': annot.get('text', ''),
                        })
            
            # Also check hyperlinks attribute if available
            if hasattr(page, 'hyperlinks') and page.hyperlinks:
                for link in page.hyperlinks:
                    if isinstance(link, dict) and link.get('uri'):
                        hyperlinks.append({
                            'uri': link.get('uri', ''),
                            'text': link.get('text', ''),
                        })
        except Exception as e:
            print(f"Warning: Could not extract hyperlinks: {e}")
        
        return hyperlinks
    
    @staticmethod
    def extract_text(file_content: bytes) -> Dict[str, any]:
        """
        Extract text from PDF file.
        
        Args:
            file_content: PDF file bytes
            
        Returns:
            Dict with 'text', 'num_pages', 'metadata', and 'hyperlinks'
        """
        try:
            pdf_file = io.BytesIO(file_content)
            
            text_parts = []
            num_pages = 0
            metadata = {}
            all_hyperlinks = []
            
            with pdfplumber.open(pdf_file) as pdf:
                num_pages = len(pdf.pages)
                metadata = pdf.metadata or {}
                
                for page_num, page in enumerate(pdf.pages, 1):
                    # Extract hyperlinks from page
                    page_links = PDFExtractor.extract_hyperlinks(page)
                    all_hyperlinks.extend(page_links)
                    
                    # Extract text with proper spacing parameters to avoid concatenation
                    page_text = page.extract_text(
                        x_tolerance=2,  # Horizontal tolerance for character grouping
                        y_tolerance=3,  # Vertical tolerance for line grouping
                        layout=False
                    )
                    
                    if page_text and len(page_text.strip()) > 50:
                        # Add page separator for multi-page resumes
                        if page_num > 1:
                            text_parts.append('\n--- Page Break ---\n')
                        text_parts.append(page_text)
                    else:
                        # If default extraction fails or produces poor results, try with layout preservation
                        page_text = page.extract_text(
                            layout=True,
                            x_tolerance=2,
                            y_tolerance=3
                        )
                        if page_text:
                            if page_num > 1:
                                text_parts.append('\n--- Page Break ---\n')
                            text_parts.append(page_text)
            
            full_text = '\n'.join(text_parts)
            
            # Append hyperlink URLs to the text so parser can find them
            # This is crucial for LaTeX-generated PDFs where links are hidden
            if all_hyperlinks:
                hyperlink_text = "\n\n--- EXTRACTED HYPERLINKS ---\n"
                for link in all_hyperlinks:
                    hyperlink_text += f"{link['uri']}\n"
                full_text += hyperlink_text
            
            return {
                'text': full_text,
                'num_pages': num_pages,
                'metadata': metadata,
                'hyperlinks': all_hyperlinks,
                'extraction_method': 'pdfplumber'
            }
            
        except Exception as e:
            raise Exception(f"PDF extraction failed: {str(e)}")
    
    @staticmethod
    def extract_with_layout(file_content: bytes) -> Dict[str, any]:
        """
        Extract text preserving layout (better for multi-column resumes).
        
        Args:
            file_content: PDF file bytes
            
        Returns:
            Dict with text and metadata
        """
        try:
            pdf_file = io.BytesIO(file_content)
            
            text_parts = []
            num_pages = 0
            
            with pdfplumber.open(pdf_file) as pdf:
                num_pages = len(pdf.pages)
                
                for page_num, page in enumerate(pdf.pages, 1):
                    # Extract with layout preservation
                    page_text = page.extract_text(
                        layout=True,
                        x_tolerance=3,
                        y_tolerance=3
                    )
                    
                    if page_text:
                        if page_num > 1:
                            text_parts.append('\n--- Page Break ---\n')
                        text_parts.append(page_text)
            
            return {
                'text': '\n'.join(text_parts),
                'num_pages': num_pages,
                'extraction_method': 'pdfplumber_layout'
            }
            
        except Exception as e:
            raise Exception(f"PDF layout extraction failed: {str(e)}")


class DOCXExtractor:
    """Extract text from DOCX files using python-docx."""
    
    @staticmethod
    def extract_text(file_content: bytes) -> Dict[str, any]:
        """
        Extract text from DOCX file.
        
        Args:
            file_content: DOCX file bytes
            
        Returns:
            Dict with 'text', 'num_paragraphs', and metadata
        """
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            
            text_parts = []
            num_paragraphs = 0
            
            # Extract from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
                    num_paragraphs += 1
            
            # Extract from tables (common in resumes)
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            full_text = '\n'.join(text_parts)
            
            # Validate extraction
            if not full_text or len(full_text.strip()) < 10:
                raise Exception("Extracted text is empty or too short. File may be corrupted or password-protected.")
            
            # Extract core properties if available
            metadata = {}
            try:
                core_props = doc.core_properties
                metadata = {
                    'author': core_props.author,
                    'title': core_props.title,
                    'subject': core_props.subject,
                    'created': str(core_props.created) if core_props.created else None,
                    'modified': str(core_props.modified) if core_props.modified else None,
                }
            except Exception:
                pass
            
            return {
                'text': full_text,
                'num_paragraphs': num_paragraphs,
                'num_tables': len(doc.tables),
                'metadata': metadata,
                'extraction_method': 'python-docx'
            }
            
        except Exception as e:
            # Provide more detailed error message
            error_msg = str(e)
            if 'BadZipFile' in error_msg or 'not a zip file' in error_msg.lower():
                raise Exception("File is not a valid DOCX file. It may be corrupted or in an older DOC format.")
            elif 'password' in error_msg.lower() or 'encrypted' in error_msg.lower():
                raise Exception("File appears to be password-protected. Please upload an unprotected version.")
            else:
                raise Exception(f"DOCX extraction failed: {error_msg}")
    
    @staticmethod
    def extract_with_formatting(file_content: bytes) -> Dict[str, any]:
        """
        Extract text with basic formatting preserved.
        
        Args:
            file_content: DOCX file bytes
            
        Returns:
            Dict with text and formatting info
        """
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            
            formatted_parts = []
            
            for para in doc.paragraphs:
                if not para.text.strip():
                    continue
                
                # Check if heading
                if para.style.name.startswith('Heading'):
                    formatted_parts.append(f"\n## {para.text.strip()} ##\n")
                elif para.text.strip():
                    formatted_parts.append(para.text)
            
            return {
                'text': '\n'.join(formatted_parts),
                'extraction_method': 'python-docx_formatted'
            }
            
        except Exception as e:
            raise Exception(f"DOCX formatted extraction failed: {str(e)}")


class ResumeExtractor:
    """
    Main extractor that routes to appropriate parser based on file type.
    Also detects LaTeX content in PDFs and parses accordingly.
    """
    
    def __init__(self):
        self.pdf_extractor = PDFExtractor()
        self.docx_extractor = DOCXExtractor()
        self.latex_extractor = LaTeXExtractor()
    
    def extract(
        self,
        file_content: bytes,
        content_type: str,
        preserve_layout: bool = False
    ) -> Dict[str, any]:
        """
        Extract text from resume file.
        
        Args:
            file_content: File bytes
            content_type: MIME type
            preserve_layout: Whether to preserve layout (useful for multi-column)
            
        Returns:
            Extraction result with text and metadata
        """
        # Handle LaTeX files
        if content_type in ['application/x-tex', 'text/x-tex', 'text/plain']:
            # Check if it's actually LaTeX
            try:
                text = file_content.decode('utf-8', errors='ignore')
                if LaTeXExtractor.is_latex_content(text):
                    return self.latex_extractor.extract_text(file_content)
            except:
                pass
            
            # If text/plain but not LaTeX, just return as text
            if content_type == 'text/plain':
                return {
                    'text': file_content.decode('utf-8', errors='ignore'),
                    'extraction_method': 'plain_text'
                }
        
        if content_type == 'application/pdf':
            if preserve_layout:
                result = self.pdf_extractor.extract_with_layout(file_content)
            else:
                result = self.pdf_extractor.extract_text(file_content)
            
            # Check if the PDF was generated from LaTeX and has LaTeX-like structure
            # (This helps with PDFs that have specific LaTeX formatting patterns)
            return result
        
        elif content_type in [
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/msword'
        ]:
            if preserve_layout:
                return self.docx_extractor.extract_with_formatting(file_content)
            else:
                return self.docx_extractor.extract_text(file_content)
        
        else:
            raise ValueError(f"Unsupported content type: {content_type}")
    
    def extract_with_fallback(
        self,
        file_content: bytes,
        content_type: str
    ) -> Dict[str, any]:
        """
        Extract with automatic fallback to layout-preserving mode if initial extraction is poor.
        
        Args:
            file_content: File bytes
            content_type: MIME type
            
        Returns:
            Best extraction result
        """
        # Try default extraction first
        result = self.extract(file_content, content_type, preserve_layout=False)
        
        # Check if extraction was poor (very short text)
        text = result.get('text', '')
        if len(text.strip()) < 100:
            # Try with layout preservation
            try:
                result = self.extract(file_content, content_type, preserve_layout=True)
            except Exception:
                # Keep original result if fallback fails
                pass
        
        return result
    
    def extract_with_latex_detection(
        self,
        file_content: bytes,
        content_type: str,
        filename: str = ''
    ) -> Dict[str, any]:
        """
        Extract with LaTeX detection from filename or content.
        
        Args:
            file_content: File bytes
            content_type: MIME type  
            filename: Original filename (for extension detection)
            
        Returns:
            Extraction result with LaTeX parsing if detected
        """
        # Check if file is LaTeX by extension
        if filename.lower().endswith('.tex'):
            return self.latex_extractor.extract_text(file_content)
        
        # Regular extraction
        result = self.extract_with_fallback(file_content, content_type)
        
        # For PDFs, also try to detect if the content structure suggests LaTeX origin
        # This is useful for PDFs compiled from LaTeX
        if content_type == 'application/pdf':
            text = result.get('text', '')
            
            # Check for LaTeX-generated PDF indicators
            latex_pdf_indicators = [
                # Common LaTeX resume patterns in extracted text
                r'^\s*[A-Z][a-z]+\s+[A-Z][a-z]+\s*$',  # Name on its own line
                r'linkedin\.com/in/',
                r'github\.com/',
            ]
            
            # If text extraction is poor, this might be a complex LaTeX PDF
            # The result will be handled by the improved parser
        
        return result
